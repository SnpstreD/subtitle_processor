import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


TURQUOISE_NAMES = ['МЭТТ', 'ТАЛЕСИН', 'ЛИАМ', 'СЭМ', 'ТРЭВИС']
RED_NAMES = ['ЛОРА', 'МАРИША', 'ЭШЛИ']

CHAR_COLORS = {
    'МЭТТ': WD_COLOR_INDEX.YELLOW,
    'ТАЛЕСИН': WD_COLOR_INDEX.VIOLET,
    'ЛИАМ': WD_COLOR_INDEX.BLUE,
    'СЭМ': WD_COLOR_INDEX.BRIGHT_GREEN,
    'ТРЭВИС': WD_COLOR_INDEX.GREEN,

    'ЛОРА': WD_COLOR_INDEX.PINK,
    'МАРИША': WD_COLOR_INDEX.RED,
    'ЭШЛИ': WD_COLOR_INDEX.TURQUOISE,
}

DEFAULT_COLOR = WD_COLOR_INDEX.GRAY_25


def parse_srt(file_name):
    with open(file_name, 'r', encoding='utf-8') as f:
        blocks = f.read().split('\n\n')

    lines = []
    for block in blocks:
        parts = block.split('\n')
        if len(parts) >= 3:
            timecode = parts[1]
            timecode_mod = [x.split(',')[0] for x in timecode.split(' --> ')]
            text = ' '.join(parts[2:])
            lines.append(f"{timecode_mod[0]} - {timecode_mod[1]} - {text}")

    merged = []
    current = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        match = re.match(
            r'(\d{2}:\d{2}:\d{2}) - \d{2}:\d{2}:\d{2} - ([А-ЯЁ]+):\s*(.*)',
            line
        )

        if match:
            if current:
                merged.append(current)

            current = {
                'time': match.group(1),
                'person': match.group(2),
                'text': match.group(3),
                'color': CHAR_COLORS.get(match.group(2), DEFAULT_COLOR)
            }

        elif current:
            text_part = re.sub(
                r'^\d{2}:\d{2}:\d{2} - \d{2}:\d{2}:\d{2} - ', '',
                line
            )
            current['text'] += ' ' + text_part

    if current:
        merged.append(current)

    return merged


def create_docx(merged_text, output_file_name):

    doc = Document()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Таймкод'
    hdr_cells[1].text = 'Персонаж'
    hdr_cells[2].text = 'Текст'

    for entry in merged_text:
        row_cells = table.add_row().cells

        row_cells[0].text = entry['time']

        person_run = row_cells[1].paragraphs[0].add_run(entry['person'])
        person_run.font.highlight_color = entry['color']
        person_run.font.bold = True

        row_cells[2].text = entry['text']

    doc.save(output_file_name)


if __name__ == '__main__':

    subs_file = 'example.srt'
    output_file = 'example.docx'

    subs_merged = parse_srt(subs_file)
    create_docx(subs_merged, output_file)

    print(f'✅ {output_file} сохранен.')
